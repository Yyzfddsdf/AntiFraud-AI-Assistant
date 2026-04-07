#!/usr/bin/env python
# -*- coding: utf-8 -*-

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = REPO_ROOT / "competition_docs" / "专利" / "源码文档"
OUTPUT_FILE = OUTPUT_DIR / "反诈卫士（Sentinel AI）_源程序文档_V1.0.docx"

SOFTWARE_NAME = "反诈卫士（Sentinel AI）"
SOFTWARE_VERSION = "V1.0"
LINES_PER_PAGE = 50
PAGES_PER_SECTION = 30
TOTAL_PAGES = PAGES_PER_SECTION * 2
TOTAL_LINES = LINES_PER_PAGE * TOTAL_PAGES
CODE_FONT_SIZE = 6
ROW_HEIGHT_MM = 3.15


@dataclass(frozen=True)
class Segment:
    path: str
    start: int = 1
    end: int | None = None


FRONT_SEGMENTS: tuple[Segment, ...] = (
    Segment("cmd/api/main.go"),
    Segment("internal/bootstrap/server/server.go"),
    Segment("internal/modules/login/adapters/inbound/http/controllers/auth.go"),
    Segment("internal/modules/login/adapters/inbound/http/controllers/auth_handler.go"),
    Segment("internal/modules/login/adapters/inbound/http/controllers/auth_service.go"),
    Segment("internal/modules/login/adapters/inbound/http/middleware/auth_middleware.go"),
    Segment("internal/modules/chat/application/usecase.go"),
    Segment("internal/modules/chat/adapters/inbound/http/chat_handler.go"),
    Segment("frontend/desktop-vue/src/main.js"),
    Segment("frontend/desktop-vue/src/App.vue"),
    Segment("frontend/desktop-vue/src/modules/chat/useChatModule.js"),
    Segment("frontend/desktop-vue/src/app/useDesktopApp.js"),
)

BACK_SEGMENTS: tuple[Segment, ...] = (
    Segment("internal/modules/multi_agent/core/main_agent.go"),
    Segment("internal/modules/multi_agent/adapters/outbound/tool/risk_assessment_tool.go"),
    Segment("frontend/desktop-vue/src/modules/geo/useGeoRiskMapModule.js"),
    Segment("Android/app/src/main/java/com/example/myapplication/SentinelRepository.kt"),
    Segment("Android/app/src/main/java/com/example/myapplication/QuickAnalyzeOverlayService.kt"),
)


def _set_cell_margins(cell, top: int = 0, start: int = 40, bottom: int = 0, end: int = 40) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders_none(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def _add_field(paragraph, field_name: str, default_text: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = default_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(8)
    section.bottom_margin = Mm(8)
    section.left_margin = Mm(8)
    section.right_margin = Mm(8)
    section.header_distance = Mm(5)
    section.footer_distance = Mm(5)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Consolas"
    normal_style.font.size = Pt(CODE_FONT_SIZE)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after = Pt(0)

    header_run = header_para.add_run(f"{SOFTWARE_NAME} {SOFTWARE_VERSION}")
    header_run.font.name = "SimSun"
    header_run.font.size = Pt(10.5)
    header_run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(0)
    footer_para.paragraph_format.space_after = Pt(0)

    prefix = footer_para.add_run("第 ")
    prefix.font.name = "SimSun"
    prefix.font.size = Pt(10.5)
    prefix._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    _add_field(footer_para, "PAGE")
    suffix = footer_para.add_run(" 页")
    suffix.font.name = "SimSun"
    suffix.font.size = Pt(10.5)
    suffix._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def _read_segment_lines(segment: Segment) -> list[str]:
    path = REPO_ROOT / segment.path
    if not path.exists():
        raise FileNotFoundError(f"missing source file: {segment.path}")
    content = path.read_text(encoding="utf-8").splitlines()
    start_index = max(segment.start - 1, 0)
    end_index = segment.end if segment.end is not None else len(content)
    lines = content[start_index:end_index]
    if not lines:
        raise ValueError(f"empty segment: {segment.path}")
    output = [f"/* FILE: {segment.path} */"]
    for offset, line in enumerate(lines, start=segment.start):
        normalized_line = line.replace("\t", "    ")
        output.append(f"{offset:04d}  {normalized_line}")
    return output


def _collect_section_lines(segments: Iterable[Segment], target_count: int) -> list[str]:
    collected: list[str] = []
    for segment in segments:
        for line in _read_segment_lines(segment):
            if len(collected) >= target_count:
                return collected
            collected.append(line)
    if len(collected) < target_count:
        raise ValueError(f"insufficient source lines: need {target_count}, got {len(collected)}")
    return collected


def _page_chunks(lines: list[str], size: int) -> list[list[str]]:
    return [lines[index:index + size] for index in range(0, len(lines), size)]


def _append_page(document: Document, page_lines: list[str]) -> None:
    table = document.add_table(rows=LINES_PER_PAGE, cols=1)
    table.autofit = False
    _set_table_borders_none(table)
    for row_index in range(LINES_PER_PAGE):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Mm(ROW_HEIGHT_MM)
        cell = row.cells[0]
        _set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(CODE_FONT_SIZE)
        run = paragraph.add_run(page_lines[row_index] if row_index < len(page_lines) else "")
        run.font.name = "Consolas"
        run.font.size = Pt(CODE_FONT_SIZE)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def build_document() -> Path:
    front_lines = _collect_section_lines(FRONT_SEGMENTS, LINES_PER_PAGE * PAGES_PER_SECTION)
    back_lines = _collect_section_lines(BACK_SEGMENTS, LINES_PER_PAGE * PAGES_PER_SECTION)
    all_lines = front_lines + back_lines
    if len(all_lines) != TOTAL_LINES:
        raise ValueError(f"unexpected line count: {len(all_lines)} != {TOTAL_LINES}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure_document(document)
    for page_index, chunk in enumerate(_page_chunks(all_lines, LINES_PER_PAGE), start=1):
        _append_page(document, chunk)
        if page_index < TOTAL_PAGES:
            document.add_page_break()
    document.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    output_path = build_document()
    print(str(output_path))

