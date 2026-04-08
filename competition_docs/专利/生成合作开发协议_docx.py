from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DEFAULT_OUTPUT = Path(r"competition_docs/专利/反诈卫士（Sentinel AI）V1.0_合作开发协议_可签字版.docx")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a signature-ready cooperation agreement DOCX.")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT), help="Output .docx path.")
    parser.add_argument("--software-name", default="反诈卫士（Sentinel AI）", help="Software name.")
    parser.add_argument("--version", default="V1.0", help="Software version.")
    parser.add_argument("--party-a-name", default="杨艺哲", help="Party A name.")
    parser.add_argument("--party-a-id", default="330501200506237838", help="Party A ID number.")
    parser.add_argument("--party-a-share", default="60%", help="Party A share.")
    parser.add_argument("--party-b-name", default="祝晟杰", help="Party B name.")
    parser.add_argument("--party-b-id", default="330182200603310212", help="Party B ID number.")
    parser.add_argument("--party-b-share", default="40%", help="Party B share.")
    parser.add_argument(
        "--representative",
        default="甲方 杨艺哲",
        help="Person responsible for registration filing and follow-up coordination.",
    )
    return parser.parse_args()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in ("top", top), ("start", start), ("bottom", bottom), ("end", end):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, *, name="宋体", size=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_paragraph(document, text="", *, align=WD_ALIGN_PARAGRAPH.LEFT, size=11, bold=False, space_before=0, space_after=0):
    p = document.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = 1.3
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "A6A6A6")


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def fill_cell(cell, text, *, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, shade=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if shade:
        set_cell_shading(cell, shade)


def add_key_value_table(document, rows):
    table = document.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    style_table(table)
    widths = [Cm(3.0), Cm(5.0), Cm(3.4), Cm(5.6)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    for row_idx, (k1, v1, k2, v2) in enumerate(rows):
        fill_cell(table.rows[row_idx].cells[0], k1, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade="F3F6F8")
        fill_cell(table.rows[row_idx].cells[1], v1)
        fill_cell(table.rows[row_idx].cells[2], k2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade="F3F6F8")
        fill_cell(table.rows[row_idx].cells[3], v2)
    return table


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)


def add_signature_table(document, a_name, b_name, a_id, b_id, a_share, b_share):
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    style_table(table)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(8.0)
    headers = (
        f"甲方：{a_name}  |  身份证号：{a_id}  |  份额：{a_share}",
        f"乙方：{b_name}  |  身份证号：{b_id}  |  份额：{b_share}",
    )
    fill_cell(table.rows[0].cells[0], headers[0], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade="EAF2F8")
    fill_cell(table.rows[0].cells[1], headers[1], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade="EAF2F8")
    fill_cell(table.rows[1].cells[0], "签字：________________________", size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell(table.rows[1].cells[1], "签字：________________________", size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell(table.rows[2].cells[0], "日期：______年____月____日", size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell(table.rows[2].cells[1], "日期：______年____月____日", size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    return table


def build_document(output: Path, software_name: str, version: str, a_name: str, a_id: str, a_share: str, b_name: str, b_id: str, b_share: str, representative: str):
    doc = Document()
    set_document_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("合作开发协议")
    set_run_font(title_run, name="黑体", size=16, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run(f"软件名称：{software_name}    版本号：{version}")
    set_run_font(subtitle_run, size=11, bold=True)

    add_paragraph(
        doc,
        f"甲方：{a_name}（身份证号：{a_id}）\n乙方：{b_name}（身份证号：{b_id}）\n\n甲乙双方经友好协商，就共同开发上述软件事宜达成如下协议。",
        size=11,
        space_after=6,
    )

    add_paragraph(doc, "一、合作方式与权利份额", size=12, bold=True, space_before=4, space_after=4)
    add_paragraph(doc, f"1. 甲乙双方共同参与软件开发，各自承担相应开发工作，共同享有软件著作权。", size=11)
    add_paragraph(doc, f"2. 软件著作权按份共有：甲方占 {a_share}，乙方占 {b_share}。", size=11)
    add_paragraph(doc, f"3. 软件名称、版本号及权利人信息应与软件著作权登记材料保持一致。", size=11)

    add_paragraph(doc, "二、软件项目基本信息", size=12, bold=True, space_before=8, space_after=4)
    add_paragraph(doc, f"1. 软件名称：{software_name}", size=11)
    add_paragraph(doc, f"2. 版本号：{version}", size=11)
    add_paragraph(doc, "3. 开发完成日期：2026 年 4 月 1 日", size=11)
    add_paragraph(doc, "4. 软件类型：AI反诈类安全工具软件", size=11)
    add_paragraph(doc, "5. 主要功能：反诈预警、风险识别、辅助拦截与安全提示等功能。", size=11)

    add_paragraph(doc, "三、著作权归属与行使", size=12, bold=True, space_before=8, space_after=4)
    add_paragraph(doc, "1. 软件著作权归甲乙双方按份共有，双方对共有著作权行使需协商一致。", size=11)
    add_paragraph(doc, "2. 软件的使用、许可、转让、收益分配等事项，按本协议约定的份额比例执行。", size=11)
    add_paragraph(doc, "3. 未经对方书面同意，任何一方不得将共有著作权转让给第三方或设立共有人优先权以外的限制。", size=11)

    add_paragraph(doc, "四、登记办理安排", size=12, bold=True, space_before=8, space_after=4)
    add_paragraph(doc, f"1. 由 {representative} 作为登记办理代表，负责线上填报、材料提交及证书领取等相关事宜。", size=11)
    add_paragraph(doc, "2. 甲乙双方均应按登记要求及时完成实名认证、签字确认及必要材料的提供，相互配合完成登记办理。", size=11)

    add_paragraph(doc, "五、保密义务", size=12, bold=True, space_before=8, space_after=4)
    add_paragraph(doc, "1. 双方对开发过程中获悉的源代码、技术方案及相关未公开信息负有保密义务。", size=11)
    add_paragraph(doc, "2. 本协议终止后，保密义务仍继续有效，直至相关信息依法公开或获得对方书面解除。", size=11)

    add_paragraph(doc, "六、争议解决", size=12, bold=True, space_before=8, space_after=4)
    add_paragraph(doc, "1. 因本协议发生争议的，双方应先行友好协商。", size=11)
    add_paragraph(doc, "2. 协商不成的，可向有管辖权的人民法院提起诉讼。", size=11)

    add_paragraph(doc, "七、其他", size=12, bold=True, space_before=8, space_after=4)
    add_paragraph(doc, "1. 本协议自双方签字之日起生效。", size=11)
    add_paragraph(doc, "2. 本协议一式两份，甲乙双方各执一份，具有同等法律效力。", size=11)

    add_paragraph(doc, "八、签署页", size=12, bold=True, space_before=10, space_after=4)
    add_signature_table(doc, a_name, b_name, a_id, b_id, a_share, b_share)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> None:
    args = parse_args()
    output = Path(args.output)
    build_document(
        output=output,
        software_name=args.software_name,
        version=args.version,
        a_name=args.party_a_name,
        a_id=args.party_a_id,
        a_share=args.party_a_share,
        b_name=args.party_b_name,
        b_id=args.party_b_id,
        b_share=args.party_b_share,
        representative=args.representative,
    )
    print(f"Generated: {output.resolve()}")


if __name__ == "__main__":
    main()
